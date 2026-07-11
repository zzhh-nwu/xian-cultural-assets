"""
生成大赛提交文档：赛题方案文档 + 元数据标准与知识图谱设计文档
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '8B0000')

    # 数据行
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            if c == 0:
                run = p.add_run(str(cell_text))
                run.bold = True
            else:
                run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            if r % 2 == 0:
                set_cell_shading(cell, 'F8F8F8')

    return table


def gen_solution_doc():
    """生成赛题方案文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 设置标题样式
    for i, (style_id, font_name, size, color) in enumerate([
        ('Heading 1', 'SimHei', 22, '8B0000'),
        ('Heading 2', 'SimHei', 16, 'B22222'),
        ('Heading 3', 'SimHei', 13, '333333'),
    ]):
        s = doc.styles[style_id]
        s.font.name = font_name
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor(*[int(color[j:j+2], 16) for j in range(0, 6, 2)])
        s.font.bold = True
        s.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # ============ 封面 ============
    for _ in range(6):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('第三届大学生数据要素素质大赛')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('赛题方案文档')
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xB2, 0x22, 0x22)
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('陕西文物/非遗/文旅数字资产生成与内容创作')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    for _ in range(4):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 30)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    doc.add_paragraph('')

    for text in ['参赛团队：[团队名称]', '提交日期：2025年7月']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    doc.add_page_break()

    # ============ 一、项目背景与需求分析 ============
    doc.add_heading('一、项目背景与需求分析', level=1)

    doc.add_heading('1.1 项目背景', level=2)
    doc.add_paragraph(
        '陕西作为中华民族和华夏文明的重要发祥地之一，拥有极为丰富的文物、非物质文化遗产和文化旅游资源。'
        '西安作为十三朝古都，是中华文明的重要象征。然而，当前文化遗产领域面临三大核心挑战：'
    )

    items_1_1 = [
        ('数字化程度不均：', '部分知名文化遗产已完成数字化采集，但大量二、三级文物保护单位和非遗项目的数字资产化程度仍然较低，信息以非结构化形式分散存储。'),
        ('标准不统一：', '各类文化数据的格式、元数据标准、分类体系各异，导致数据无法互通共享，难以形成统一的文化数字资产体系。'),
        ('传播方式单一：', '传统文化传播主要依赖线下参观和简单图文介绍，缺乏利用AI技术进行创新内容创作和传播的有效路径。'),
    ]
    for title, desc in items_1_1:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p.add_run(desc)

    doc.add_heading('1.2 需求分析', level=2)
    doc.add_paragraph(
        '针对上述挑战，本项目围绕文化资源数字化到数字资产化再到内容产品化的核心路径，重点解决以下需求：'
    )

    needs = [
        ('数据采集与资产化预处理：', '面向西安文物、非遗、景区公开资源开展系统化数据采集、清洗和结构化处理，形成符合数字资产登记要求的基础数据集。'),
        ('AI智能体构建：', '开发文化内容创作智能体，实现文化数据的智能分类、关键词标签提取、文化要素识别等功能，大幅提升文化资源治理效率。'),
        ('多模态内容生成：', '基于智能体自动生成文案、AI图像生成提示词、视频分镜头脚本等多种形态的创作内容，助力文化遗产的数字化传播。'),
    ]
    for title, desc in needs:
        p = doc.add_paragraph()
        run = p.add_run('• ' + title)
        run.bold = True
        p.add_run(desc)

    # ============ 二、系统架构设计 ============
    doc.add_heading('二、系统架构设计', level=1)

    doc.add_heading('2.1 整体架构', level=2)
    doc.add_paragraph(
        '本系统采用四层架构设计，自上而下分别为：展示层、智能体层、知识层和数据层。'
        '各层职责明确、松耦合，支持独立扩展和升级。'
    )

    # 架构图
    arch_layers = [
        '┌───────────────────────────────────────────────────────┐',
        '│  展示层：Streamlit Web | 数据浏览 | 智能体工作台 | 分析  │',
        '├───────────────────────────────────────────────────────┤',
        '│  智能体层：Cultural Agent | 分类 | 标注 | 提取 | 生成   │',
        '├───────────────────────────────────────────────────────┤',
        '│  知识层：Knowledge Graph | 实体识别 | 关系 | RDF语义    │',
        '├───────────────────────────────────────────────────────┤',
        '│  数据层：Data Pipeline | 采集 | 清洗 | 标准化 | 登记    │',
        '└───────────────────────────────────────────────────────┘',
    ]
    for line in arch_layers:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)

    doc.add_heading('2.2 技术栈', level=2)
    add_styled_table(doc,
        ['层次', '采用技术'],
        [
            ['数据采集与处理', 'Python 3.13 | Pandas | BeautifulSoup4 | requests'],
            ['知识图谱', 'NetworkX | PyVis | RDFlib | GraphML'],
            ['AI智能体', 'Python Pydantic | DeepSeek API / 规则引擎 | 多模态支持'],
            ['Web展示', 'Streamlit | Plotly | HTML/CSS'],
            ['文档生成', 'docx (Node.js) | Python-docx'],
        ]
    )

    # ============ 三、数据处理流程 ============
    doc.add_heading('三、数据处理流程', level=1)

    doc.add_heading('3.1 数据采集', level=2)
    doc.add_paragraph(
        '数据采集严格遵守大赛数据使用规范，所有数据来源于公开渠道，聚焦西安文物、非遗、文旅三大领域。'
        '主要数据来源包括：百度百科、中国非物质文化遗产网、西安市文化和旅游局官网、各景区官方公开信息等。'
    )

    add_styled_table(doc,
        ['数据类别', '数据量', '主要来源', '采集方式'],
        [
            ['文物数据', '20条', '百度百科/政府公开信息', '结构化提取+人工核验'],
            ['非遗数据', '18条', '中国非物质文化遗产网/地方志', '结构化提取+人工核验'],
            ['文旅数据', '16条', '文旅局官网/景区公开信息', '结构化提取+人工核验'],
            ['总计', '54条', '公开合法来源', '标准化处理'],
        ]
    )

    doc.add_heading('3.2 数据清洗与结构化', level=2)
    doc.add_paragraph('数据处理流程包括五个关键步骤：')

    steps = [
        ('文本清洗：', '去除多余空格、规范化标点符号、统一括号格式，确保文本数据的一致性。'),
        ('数据去重：', '基于名称字段进行唯一性检查，去除重复记录，保证数据集的精炼性。'),
        ('字段标准化：', '统一保护级别、非遗级别、景区等级等字段的标准化表达，建立统一的分类体系。'),
        ('质量验证：', '对每条记录进行必填字段、关键词标签的完整性检查，生成质量报告。'),
        ('资产元数据生成：', '依照DC（Dublin Core）元数据标准，为数据集生成符合数字资产登记要求的元数据描述。'),
    ]
    for title, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run('• ' + title)
        run.bold = True
        p.add_run(desc)

    doc.add_heading('3.3 数据集质量', level=2)
    doc.add_paragraph('经过完整的清洗和结构化处理，数据集质量指标如下：')

    quality_items = [
        '总记录数：54条（文物20条、非遗18条、文旅16条）',
        '数据完整度：平均85%以上',
        '质量问题数：0（全部通过质量验证）',
        '格式支持：JSON、CSV、Excel三种格式',
        '每条记录包含：ID、名称、类型、类别、年代/级别、关键词标签、结构化扩展数据',
    ]
    for item in quality_items:
        doc.add_paragraph('• ' + item)

    # ============ 四、AI智能体设计 ============
    doc.add_heading('四、AI智能体设计', level=1)

    doc.add_heading('4.1 智能体定位', level=2)
    doc.add_paragraph(
        '文化内容创作智能体（Cultural Content Agent）是本项目的核心引擎，'
        '定位为文化数据治理与内容创作的双功能智能体，能够实现：'
    )

    agent_functions = [
        '对文化遗产数据进行智能分类和关键词标注',
        '从文本中提取文化要素（人物、地点、事件、朝代、技艺、符号）',
        '自动生成多模态创作内容（文案、图像提示词、视频脚本）',
        '进行内容合规性检查，确保符合文博/非遗传播规范',
    ]
    for f in agent_functions:
        doc.add_paragraph('• ' + f)

    doc.add_heading('4.2 技术实现', level=2)
    doc.add_paragraph(
        '智能体采用AI模型 + 规则引擎的双模式设计，确保在不同环境下的可用性：'
    )

    impl_items = [
        ('AI模式：', '通过DeepSeek API实现深度语义理解和高质量内容生成。支持文化分类、要素提取、多平台内容创作等复杂任务。System Prompt经过精心设计，内嵌文博/非遗传播规范。'),
        ('规则模式：', '当API不可用时，智能体自动切换到基于规则的生成模式。内置丰富的西安文化知识库（关键词库、朝代信息、人物关系、地点描述等），确保离线环境下仍能产出高质量内容。'),
        ('输出结构化：', '所有输出均采用Pydantic数据模型进行结构化约束，确保生成的每一条内容都包含完整的元信息（类型、标签、目标平台、字数等）。'),
    ]
    for title, desc in impl_items:
        p = doc.add_paragraph()
        run = p.add_run('• ' + title)
        run.bold = True
        p.add_run(desc)

    doc.add_heading('4.3 核心功能模块', level=2)
    add_styled_table(doc,
        ['功能模块', '功能描述', '技术特点'],
        [
            ['智能分类标注', '对文本进行文物/非遗/文旅分类，生成关键词标签', 'AI语义理解+关键词匹配双引擎'],
            ['文化要素提取', '识别文本中的人物、地点、朝代、技艺等文化要素', '命名实体识别+文化知识库'],
            ['文案生成', '生成适合小红书/抖音/微信等平台的推广文案', '多平台风格适配+字数控制'],
            ['图像提示词生成', '输出中英文AI图像生成提示词', '指定构图/风格/色彩/氛围'],
            ['视频脚本生成', '输出分镜头脚本（镜头号/画面/旁白/时长）', '行业标准脚本格式'],
            ['合规性检查', '检查内容是否符合文博/非遗传播规范', '敏感词检测+历史准确性校验'],
        ]
    )

    # ============ 五、多模态内容生成方案 ============
    doc.add_heading('五、多模态内容生成方案', level=1)

    doc.add_heading('5.1 文案生成', level=2)
    for item in [
        '目标平台：小红书、抖音、微信公众号、微博、B站等多平台支持',
        '生成策略：根据平台特性自动调整文风（小红书亲切活泼、微信公众号深度长文、抖音短小精悍）',
        '内容要素：标题 + 正文 + 文化标签 + 互动引导 + 平台专属话题标签',
    ]:
        doc.add_paragraph('• ' + item)

    doc.add_heading('5.2 AI图像提示词生成', level=2)
    for item in [
        '输出格式：中英文双语提示词，适配主流AI图像生成工具（如Midjourney、Stable Diffusion、DALL-E等）',
        '风格多样性：写实摄影风、国风插画风、夜景氛围风、艺术插画风等多风格支持',
        '提示词要素：构图描述 + 风格指定 + 色彩指导 + 光线氛围 + 细节要求',
    ]:
        doc.add_paragraph('• ' + item)

    doc.add_heading('5.3 视频分镜头脚本生成', level=2)
    for item in [
        '脚本格式：行业标准分镜头格式，包含镜头号、场景描述、时长、运镜方式、画面描述、旁白/对白',
        '叙事结构：开场吸引力（5s）→ 主体内容展示（循环展示）→ 结尾升华（8s），符合短视频传播规律',
        '音频设计：每个镜头搭配背景音乐和音效建议，融入非遗项目原声元素',
    ]:
        doc.add_paragraph('• ' + item)

    doc.add_heading('5.4 已生成内容统计', level=2)
    doc.add_paragraph('系统已为8个核心主题批量生成了24条多模态内容：')

    add_styled_table(doc,
        ['内容主题', '文案', '图片提示词', '视频脚本'],
        [
            ['秦始皇兵马俑', '✓', '✓', '✓'],
            ['大雁塔与丝绸之路', '✓', '✓', '✓'],
            ['西安城墙', '✓', '✓', '✓'],
            ['大唐不夜城', '✓', '✓', '✓'],
            ['西安鼓乐', '✓', '✓', '✓'],
            ['秦腔艺术', '✓', '✓', '✓'],
            ['回民街美食', '✓', '✓', '✓'],
            ['华清宫长恨歌', '✓', '✓', '✓'],
        ]
    )

    # ============ 六、数字资产化路径 ============
    doc.add_heading('六、数字资产化路径', level=1)

    doc.add_heading('6.1 文化资源到数据到数字资产转化路径', level=2)
    doc.add_paragraph('本项目提出并实践了完整的文化资源到数据再到数字资产的转化路径：')

    doc.add_heading('第一阶段：资源数字化', level=3)
    for item in [
        '采集：从公开渠道系统采集西安文物、非遗、文旅资源信息',
        '整理：按统一分类体系和字段结构进行初步整理',
    ]:
        doc.add_paragraph('• ' + item)

    doc.add_heading('第二阶段：数据资产化', level=3)
    for item in [
        '清洗：执行文本清洗、去重、标准化处理',
        '标注：通过AI智能体进行文化要素标注和关键词标签生成',
        '关联：构建知识图谱，建立实体间的跨域关联',
        '登记：生成符合DC元数据标准的资产描述，完成数字资产登记',
    ]:
        doc.add_paragraph('• ' + item)

    doc.add_heading('第三阶段：内容产品化', level=3)
    for item in [
        '创作：智能体基于结构化数据自动生成多模态内容',
        '适配：根据目标平台特性进行内容形式和风格适配',
        '分发：通过Web平台展示，支持内容下载和二次使用',
    ]:
        doc.add_paragraph('• ' + item)

    doc.add_heading('6.2 知识图谱设计概要', level=2)
    doc.add_paragraph('构建了覆盖75个节点、77条边的西安文化知识图谱：')

    add_styled_table(doc,
        ['维度', '详情'],
        [
            ['节点类型（6种）', '文物(20) + 非遗(18) + 文旅(16) + 朝代(7) + 非遗类别(5) + 区域(9) = 75个'],
            ['关系类型（4种）', '文化关联(11) + 层级包含(54) + 时序先后(4) + 空间毗邻(8) = 77条'],
            ['导出格式', 'JSON + GraphML + RDF/Turtle + HTML交互可视化'],
            ['语义标准', 'RDF/OWL，支持SPARQL查询'],
        ]
    )

    # ============ 七、创新与应用价值 ============
    doc.add_heading('七、创新与应用价值', level=1)

    doc.add_heading('7.1 创新点', level=2)
    innovations = [
        ('双模式智能体架构：', 'AI模型+规则引擎双模式设计，确保智能体在不同环境（有/无API）下均可运行，提高了系统的鲁棒性和实用性。'),
        ('多模态统一生成框架：', '一套系统同时支持文案、图像提示词、视频脚本三种内容形态的自动生成，覆盖了当前主流内容平台的创作需求。'),
        ('文化知识图谱：', '首次为西安文物/非遗/文旅构建跨领域知识图谱，支持多维度关联查询和文化脉络发现。'),
        ('标准化资产登记：', '严格遵循DC元数据标准进行数字资产描述，为文化数据资产的登记、上架和交易奠定标准基础。'),
        ('全链路闭环：', '实现采集→处理→知识化→智能创作→展示的完整闭环，形成可复制的文化数字化方案。'),
    ]
    for i, (title, desc) in enumerate(innovations, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {title}')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('7.2 应用价值', level=2)
    values = [
        ('贴合西安文化特色：', '聚焦西安最具代表性的文物（兵马俑、大雁塔等）、非遗（鼓乐、秦腔等）、文旅资源（大唐不夜城等），充分体现地域文化特色，契合文物保护与非遗传承要求。'),
        ('符合文保/非遗传承要求：', '内建合规性检查机制，自动检测不当表述，确保所有输出内容尊重历史事实、弘扬中华优秀传统文化。'),
        ('可复制路径：', '模块化设计，三大核心模块（数据采集、知识图谱、智能体）可独立部署和复用。将西安方案迁移至其他文化名城（北京、南京、洛阳等），只需替换数据源即可快速构建。'),
        ('支撑资产登记与交易：', '数据集符合DC元数据标准，知识图谱支持RDF语义格式输出，可直接用于文化数据资产登记平台，支撑后续的授权、运营和交易。'),
    ]
    for title, desc in values:
        p = doc.add_paragraph()
        run = p.add_run('• ' + title)
        run.bold = True
        p.add_run(desc)

    doc.add_heading('7.3 对标评分标准', level=2)
    add_styled_table(doc,
        ['评分维度', '分值', '本项目达成情况'],
        [
            ['数据处理规范性', '30分',
             '公开数据来源合法(10) ✓ | 数据清洗有效(10) ✓ | 符合资产化要求(10) ✓'],
            ['方案完成度', '30分',
             'AI智能体治理/标注/提取(10) ✓ | 多模态生成完整(10) ✓ | 可用于展示/授权(10) ✓'],
            ['创新与应用价值', '25分',
             '贴合西安文化特色(10) ✓ | 可复制路径(10) ✓ | 支撑资产登记交易(5) ✓'],
            ['文档与展示', '15分',
             '流程清晰 ✓ | 作品完整 ✓ | Web平台直观展示 ✓'],
        ]
    )

    # ============ 八、成果清单 ============
    doc.add_heading('八、成果清单', level=1)

    doc.add_heading('8.1 核心成果（已全部完成）', level=2)
    doc.add_paragraph('1. 文物/非遗/文旅数字资产包（数据集）')
    doc.add_paragraph('   包含54条记录的西安文物/非遗/文旅结构化数据集，支持JSON/CSV/Excel三种格式，含知识图谱基础数据，可直接用于数字资产登记。')
    doc.add_paragraph('')
    doc.add_paragraph('2. 文化内容创作智能体')
    doc.add_paragraph('   Python实现的AI智能体，支持智能分类、文化要素提取、多模态内容生成（文案+图像提示词+视频脚本）、合规性检查。AI+规则双模式运行，支持有/无API环境。')

    doc.add_heading('8.2 配套文档', level=2)
    doc.add_paragraph('1. 赛题方案文档（本文档）—— 包含项目背景、系统架构、数据流程、智能体设计、创新价值等完整内容')
    doc.add_paragraph('2. 数据资产元数据标准与知识图谱设计文档 —— 详细描述DC元数据标准体系、知识图谱设计方法论、实体关系模型')
    doc.add_paragraph('3. 演示视频（≤3分钟）—— 含功能讲解和效果展示')

    doc.add_heading('8.3 辅助成果', level=2)
    for item in [
        'Streamlit Web展示平台：一站式数据资产浏览、知识图谱可视化、智能体工作台',
        '批量生成内容库：24条多模态内容（8个主题 × 3种类型）',
        '知识图谱多格式导出：JSON + GraphML + RDF/Turtle + HTML可视化',
        '数据清洗报告：完整记录清洗步骤和质量验证结果',
        '项目源代码：完整的Python代码，注释清晰，结构规范，可直接运行',
    ]:
        doc.add_paragraph('• ' + item)

    # 页尾
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 20 + '  END  ' + '━' * 20)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    run.bold = True

    # 保存
    path = os.path.join(OUTPUT_DIR, '赛题方案文档.docx')
    doc.save(path)
    print(f'✅ 方案文档已生成: {path}')
    return path


def gen_metadata_doc():
    """生成数据资产元数据标准与知识图谱设计文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 标题样式
    for style_id, font_name, size, color in [
        ('Heading 1', 'SimHei', 22, '8B0000'),
        ('Heading 2', 'SimHei', 16, 'B22222'),
        ('Heading 3', 'SimHei', 13, '333333'),
    ]:
        s = doc.styles[style_id]
        s.font.name = font_name
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor(*[int(color[j:j+2], 16) for j in range(0, 6, 2)])
        s.font.bold = True
        s.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # ========== 封面 ==========
    for _ in range(6):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('数据资产元数据标准\n与\n知识图谱设计文档')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('陕西文物/非遗/文旅数字资产生成与内容创作')
    run.font.size = Pt(18)
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    for _ in range(6):
        doc.add_paragraph('')

    for text in ['参赛团队：[团队名称]', '文档版本：v1.0', '日期：2025年7月']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(14)

    doc.add_page_break()

    # ========== 第一部分：元数据标准 ==========
    doc.add_heading('第一部分：数据资产元数据标准', level=1)

    doc.add_heading('1.1 标准选择与依据', level=2)
    doc.add_paragraph(
        '本数据集采用都柏林核心元数据标准（Dublin Core Metadata Element Set, DC）作为基础元数据标准，'
        '并结合文化数据资产的特性进行了扩展。DC标准是国际通用的资源描述元数据标准（ISO 15836），'
        '被广泛应用于数字图书馆、文化遗产数字化和开放数据等领域。'
    )

    doc.add_heading('1.2 DC核心元数据元素', level=2)

    add_styled_table(doc,
        ['DC元素', '中文名称', '本数据集取值', '说明'],
        [
            ['dc:title', '题名', '西安文物/非遗/文旅数字资产数据集', '数据集的正式名称'],
            ['dc:creator', '创建者', '大学生数据要素素质大赛参赛团队', '创建数据集的责任实体'],
            ['dc:date', '日期', '2025-07', '数据集创建的日期'],
            ['dc:type', '类型', 'Dataset', '资源类型，符合DCMI Type词汇表'],
            ['dc:format', '格式', 'JSON/CSV/Excel', '数据文件格式'],
            ['dc:language', '语言', 'zh-CN', '数据集内容的语言'],
            ['dc:coverage', '覆盖范围', '陕西省西安市', '数据集的空间覆盖范围'],
            ['dc:rights', '权限', '仅限大赛使用', '使用权限说明'],
            ['dc:description', '描述', '面向西安文物、非遗、文旅资源的综合性数字资产数据集', '数据集内容摘要'],
            ['dc:subject', '主题', '文物；非物质文化遗产；文化旅游', '数据集的主题关键词'],
            ['dc:publisher', '出版者', '大学生数据要素素质大赛', '发布责任实体'],
            ['dc:identifier', '标识符', 'XA-CULTURE-2025-001', '数据集的唯一标识'],
            ['dc:source', '来源', '百度百科；政府公开信息；文旅平台公开数据', '数据来源说明'],
            ['dc:relation', '关联', '西安文化知识图谱', '相关资源'],
        ]
    )

    doc.add_heading('1.3 扩展元数据元素', level=2)
    doc.add_paragraph('为满足文化数字资产登记和交易的特定需求，在DC标准基础上增加了以下扩展元素：')

    add_styled_table(doc,
        ['扩展元素', '中文名称', '本数据集取值', '用途'],
        [
            ['dcterms:spatial', '空间', '陕西省西安市', '精细化的空间描述'],
            ['dcterms:temporal', '时间', '新石器时代至当代', '时间范围描述'],
            ['dcat:theme', '主题分类', '文物, 非遗, 文旅', '数据目录主题分类'],
            ['custom:record_count', '记录数', '54', '数据集总记录数'],
            ['custom:categories', '分类统计', '文物20/非遗18/文旅16', '各分类记录数'],
            ['custom:quality', '数据质量', '合格（0错误）', '数据质量评估结果'],
            ['custom:completeness', '完整度', '>85%', '字段填充完整度'],
            ['custom:standard_version', '标准版本', 'v1.0', '元数据标准版本号'],
            ['custom:knowledge_graph', '知识图谱', '75节点/77边', '关联知识图谱信息'],
            ['custom:update_frequency', '更新频率', '按需更新', '数据更新计划'],
        ]
    )

    doc.add_heading('1.4 数字资产登记信息结构', level=2)
    doc.add_paragraph('每条数据记录的数字资产登记信息包含以下标准字段：')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('基础信息层：')
    run.bold = True
    for item in ['asset_id（资产唯一标识符）', 'asset_type（资产类型：cultural_relic / intangible_cultural_heritage / cultural_tourism）',
                 'asset_level（资产级别：国家级/省级/市县级）', 'geographic_scope（地理范围）', 'time_period（时间范围）']:
        doc.add_paragraph('  • ' + item)

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('管理信息层：')
    run.bold = True
    for item in ['data_completeness（数据完整度，0-1浮点数）', 'last_updated（最后更新日期）',
                 'standard_version（标准版本号）', 'license（使用授权）']:
        doc.add_paragraph('  • ' + item)

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('内容信息层：')
    run.bold = True
    for item in ['name（名称）', 'keywords（关键词标签集）', 'structured_data（结构化扩展数据）',
                 'cultural_elements（文化要素）', 'baidu_baike_url（参考链接）']:
        doc.add_paragraph('  • ' + item)

    doc.add_page_break()

    # ========== 第二部分：知识图谱设计 ==========
    doc.add_heading('第二部分：知识图谱设计', level=1)

    doc.add_heading('2.1 设计目标', level=2)
    for item in [
        '构建西安文物、非遗、文旅三大领域的统一知识网络',
        '实现跨领域实体关联和文化脉络发现',
        '支持多维度查询：按朝代、类别、区域、文化关联等',
        '提供机器可读的语义数据格式（RDF/OWL）',
        '支撑数字资产之间的关联发现和知识推理',
    ]:
        doc.add_paragraph('• ' + item)

    doc.add_heading('2.2 本体设计', level=2)
    doc.add_paragraph('知识图谱基于自定义的西安文化本体（Xi\'an Cultural Ontology, XACO）构建，核心类定义如下：')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('实体类（Entity Classes）：')
    run.bold = True

    add_styled_table(doc,
        ['类名', '标签', '属性', '说明'],
        [
            ['CulturalRelic', '文物', 'name, category, era, protection_level, keywords', '文物保护单位和馆藏文物'],
            ['IntangibleHeritage', '非遗', 'name, level, category, origin, keywords, inheritor', '非物质文化遗产项目'],
            ['CulturalTourism', '文旅资源', 'name, type, district, description, keywords', '文化旅游景区和特色项目'],
            ['Dynasty', '朝代', 'name, start_year, end_year, capital', '历史朝代时期'],
            ['District', '区域', 'name, area, population, cultural_features', '西安市行政区划'],
            ['ICHCategory', '非遗类别', 'name, description', '非遗项目类型分类'],
        ]
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('关系类型（Relation Types）：')
    run.bold = True

    add_styled_table(doc,
        ['关系类型', 'Domain', 'Range', '语义说明'],
        [
            ['xa:contains', '朝代/类别/区域', '具体实体', '层级包含关系'],
            ['xa:cultural_link', '任意文化实体', '任意文化实体', '基于关键词/主题的文化关联'],
            ['xa:precedes', '朝代', '朝代', '时间先后关系'],
            ['xa:adjacent', '区域', '区域', '空间毗邻关系'],
            ['xa:created_in', '文物/非遗', '朝代', '创建于某朝代'],
            ['xa:located_in', '文旅资源', '区域', '位于某区域'],
        ]
    )

    doc.add_heading('2.3 图谱规模与统计', level=2)
    doc.add_paragraph('当前知识图谱数据统计：')

    add_styled_table(doc,
        ['指标', '数值'],
        [
            ['节点总数', '75'],
            ['边总数', '77'],
            ['网络密度', '0.0139'],
            ['平均度', '2.05'],
            ['连通性', '非全连通（多子图）'],
            ['节点类型分布', '文物20 | 非遗18 | 文旅16 | 朝代7 | 非遗类别5 | 区域9'],
            ['关系类型分布', '层级包含54 | 文化关联11 | 空间毗邻8 | 时序先后4'],
        ]
    )

    doc.add_heading('2.4 语义格式与互操作', level=2)
    doc.add_paragraph('知识图谱支持以下格式导出，确保跨平台互操作性：')

    add_styled_table(doc,
        ['格式', '用途', '技术标准', '状态'],
        [
            ['JSON', '程序访问和API传输', '自定义JSON Schema', '✅ 已生成'],
            ['GraphML', '图分析工具（Gephi/Cytoscape）', 'GraphML 1.0', '✅ 已生成'],
            ['RDF/Turtle', '语义网和关联数据', 'RDF 1.1 / Turtle', '✅ 已生成'],
            ['HTML', '交互式可视化', 'PyVis / vis.js', '✅ 已生成'],
        ]
    )

    doc.add_heading('2.5 查询与使用示例', level=2)

    doc.add_heading('语义查询（SPARQL示例）', level=3)
    p = doc.add_paragraph()
    run = p.add_run('查询与"大雁塔"相关的所有文化实体：')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    sparql1 = """
PREFIX xa: <http://xian.culture/kg#>
SELECT ?entity ?relation
WHERE {
  { xa:XA-REL-002 xa:cultural_link ?entity . BIND("文化关联" AS ?relation) }
  UNION
  { ?entity xa:cultural_link xa:XA-REL-002 . BIND("文化关联(反向)" AS ?relation) }
}"""
    p = doc.add_paragraph()
    run = p.add_run(sparql1)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    doc.add_heading('API查询示例', level=3)
    p = doc.add_paragraph()
    run = p.add_run('通过Python API查询实体关联：')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    code = """
from data.knowledge_graph_builder import KnowledgeGraphBuilder
builder = KnowledgeGraphBuilder()
builder.build()
result = builder.query_related("大雁塔", max_depth=2)
for r in result["direct_relations"]:
    print(f"{r['target']} --[{r['relation']}]--> 大雁塔")
"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    doc.add_heading('2.6 知识图谱扩展规划', level=2)
    for item in [
        '短期：增加人物实体类（历史人物、传承人），丰富人物关联关系',
        '中期：引入事件本体（重大历史事件、文化活动），构建时空事件网络',
        '长期：对接外部开放数据（Wikidata、中国文化遗产知识图谱），实现数据互联',
    ]:
        doc.add_paragraph('• ' + item)

    doc.add_paragraph('')
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 20 + '  END  ' + '━' * 20)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    run.bold = True

    # 保存
    path = os.path.join(OUTPUT_DIR, '数据资产元数据标准与知识图谱设计文档.docx')
    doc.save(path)
    print(f'✅ 元数据标准文档已生成: {path}')
    return path


if __name__ == '__main__':
    print('=' * 60)
    print('  生成大赛提交文档')
    print('=' * 60)
    gen_solution_doc()
    gen_metadata_doc()
    print('\n✅ 所有文档生成完毕！')
